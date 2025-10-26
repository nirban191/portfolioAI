"""
Resume parser - PDF/DOCX extraction + AI parsing
"""

import io
import re
from typing import Dict, Any, Optional, Tuple
import streamlit as st

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    st.error("PDF libraries not installed. Run: pip install pypdf2 pdfplumber")

from utils.groq_client import get_groq_client
from utils.validators import ProfileSchema
from prompts.prompts import RESUME_PARSER_PROMPT


class ResumeParser:
    """Parse resumes from PDF/DOCX files"""

    def __init__(self):
        self.groq_client = get_groq_client()

    def extract_text_from_pdf(self, file_bytes: bytes) -> Tuple[bool, str, Optional[str]]:
        """
        Extract text from PDF file

        Returns:
            (success, extracted_text, error_message)
        """
        try:
            # Try pdfplumber first (better for complex layouts)
            pdf_file = io.BytesIO(file_bytes)
            text_parts = []

            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            if text_parts:
                full_text = "\n\n".join(text_parts)
                return True, full_text, None

            # Fallback to PyPDF2
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []

            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            if text_parts:
                full_text = "\n\n".join(text_parts)
                return True, full_text, None
            else:
                return False, "", "Could not extract text from PDF. File may be scanned/image-based."

        except Exception as e:
            return False, "", f"Error reading PDF: {str(e)}"

    def extract_text_from_docx(self, file_bytes: bytes) -> Tuple[bool, str, Optional[str]]:
        """
        Extract text from DOCX file

        Returns:
            (success, extracted_text, error_message)
        """
        try:
            from docx import Document

            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            if text_parts:
                full_text = "\n".join(text_parts)
                return True, full_text, None
            else:
                return False, "", "DOCX file appears to be empty"

        except Exception as e:
            return False, "", f"Error reading DOCX: {str(e)}"

    def clean_text(self, text: str) -> str:
        """Clean extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove common artifacts
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\xff]', '', text)

        return text.strip()

    def estimate_parsing_confidence(self, text: str, parsed_data: Dict) -> float:
        """
        Estimate confidence in parsing accuracy

        Returns:
            Confidence score 0.0 to 1.0
        """
        score = 0.0

        # Check for required fields
        if parsed_data.get("name"):
            score += 0.25
        if parsed_data.get("email"):
            score += 0.25
        if parsed_data.get("work_history") and len(parsed_data["work_history"]) > 0:
            score += 0.25
        if parsed_data.get("skills") and len(parsed_data["skills"]) > 0:
            score += 0.15
        if parsed_data.get("education") and len(parsed_data["education"]) > 0:
            score += 0.10

        # Penalize if text is too short (likely parsing error)
        if len(text) < 200:
            score *= 0.5

        return min(score, 1.0)

    def parse_with_ai(self, resume_text: str) -> Tuple[bool, Optional[Dict], Optional[str]]:
        """
        Parse resume text using Groq AI

        Returns:
            (success, parsed_data, error_message)
        """
        try:
            # Truncate if too long
            if len(resume_text) > 10000:
                resume_text = resume_text[:10000] + "\n\n[Truncated for length]"

            # Call Groq API
            parsed_json = self.groq_client.call_api_json(
                system_prompt=RESUME_PARSER_PROMPT,
                user_prompt=f"Resume text:\n\n{resume_text}",
                model="8b",  # Fast model for parsing
                temperature=0.2,  # Low temperature for consistency
                max_tokens=2048
            )

            if not parsed_json:
                return False, None, "AI parsing failed - no response"

            # Validate with Pydantic
            try:
                profile = ProfileSchema(**parsed_json)
                validated_data = profile.model_dump()
                return True, validated_data, None
            except Exception as validation_error:
                # Return partial data even if validation fails
                st.warning(f"Validation warning: {validation_error}")
                return True, parsed_json, None

        except Exception as e:
            return False, None, f"Error during AI parsing: {str(e)}"

    def parse_resume(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Main parsing function

        Returns:
            {
                "success": bool,
                "profile_data": dict or None,
                "raw_text": str,
                "confidence": float,
                "error": str or None
            }
        """
        # Extract text based on file type
        file_ext = filename.lower().split('.')[-1]

        if file_ext == 'pdf':
            success, raw_text, error = self.extract_text_from_pdf(file_bytes)
        elif file_ext in ['docx', 'doc']:
            success, raw_text, error = self.extract_text_from_docx(file_bytes)
        else:
            return {
                "success": False,
                "profile_data": None,
                "raw_text": "",
                "confidence": 0.0,
                "error": f"Unsupported file type: {file_ext}"
            }

        if not success:
            return {
                "success": False,
                "profile_data": None,
                "raw_text": "",
                "confidence": 0.0,
                "error": error
            }

        # Clean text
        cleaned_text = self.clean_text(raw_text)

        if len(cleaned_text) < 100:
            return {
                "success": False,
                "profile_data": None,
                "raw_text": cleaned_text,
                "confidence": 0.0,
                "error": "Resume text too short. File may be empty or unreadable."
            }

        # Parse with AI
        success, parsed_data, error = self.parse_with_ai(cleaned_text)

        if not success:
            return {
                "success": False,
                "profile_data": None,
                "raw_text": cleaned_text,
                "confidence": 0.0,
                "error": error
            }

        # Estimate confidence
        confidence = self.estimate_parsing_confidence(cleaned_text, parsed_data)

        # Add original text to parsed data
        if parsed_data:
            parsed_data["original_resume_text"] = cleaned_text
            parsed_data["parsing_confidence"] = confidence
            parsed_data["parsing_method"] = "pdf" if file_ext == "pdf" else "docx"

        return {
            "success": True,
            "profile_data": parsed_data,
            "raw_text": cleaned_text,
            "confidence": confidence,
            "error": None
        }


# Singleton instance
_resume_parser = None


def get_resume_parser() -> ResumeParser:
    """Get or create resume parser singleton"""
    global _resume_parser
    if _resume_parser is None:
        _resume_parser = ResumeParser()
    return _resume_parser
