"""
CV (Curriculum Vitae) generator - Create comprehensive academic/professional CVs
"""

import io
from typing import Dict, Any, Optional, Tuple
from datetime import datetime
import streamlit as st

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CVGenerator:
    """Generate comprehensive CV (Curriculum Vitae) in PDF and DOCX formats"""

    def __init__(self):
        pass

    def format_date_range(self, dates: str) -> str:
        """Format date range for consistency"""
        if not dates:
            return ""
        dates = dates.strip()
        dates = dates.replace(" - ", " to ")
        return dates

    def generate_pdf(self, profile_data: Dict[str, Any]) -> Tuple[bool, Optional[bytes], Optional[str]]:
        """
        Generate comprehensive CV in PDF format

        Returns:
            (success, pdf_bytes, error_message)
        """
        try:
            # Create PDF buffer
            buffer = io.BytesIO()

            # Create PDF document (no page limit for CV)
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            # Styles
            styles = getSampleStyleSheet()

            # Custom styles for CV
            style_name = ParagraphStyle(
                'CVName',
                parent=styles['Heading1'],
                fontSize=20,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=6,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )

            style_contact = ParagraphStyle(
                'CVContact',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=16,
                alignment=TA_CENTER
            )

            style_section = ParagraphStyle(
                'CVSection',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#0A2540'),
                spaceAfter=8,
                spaceBefore=12,
                fontName='Helvetica-Bold',
                borderWidth=1,
                borderColor=colors.HexColor('#667eea'),
                borderPadding=4
            )

            style_subsection = ParagraphStyle(
                'CVSubsection',
                parent=styles['Heading3'],
                fontSize=12,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=4,
                spaceBefore=6,
                fontName='Helvetica-Bold'
            )

            style_body = ParagraphStyle(
                'CVBody',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=6,
                leading=14
            )

            style_date = ParagraphStyle(
                'CVDate',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#666666'),
                spaceAfter=4,
                fontName='Helvetica-Oblique'
            )

            # Build content
            content = []

            # Header - Name
            name = profile_data.get('name', 'Your Name')
            content.append(Paragraph(name, style_name))

            # Contact Information
            contact_parts = []
            if profile_data.get('email'):
                contact_parts.append(profile_data['email'])
            if profile_data.get('phone'):
                contact_parts.append(profile_data['phone'])
            if profile_data.get('linkedin_url'):
                contact_parts.append(f"LinkedIn: {profile_data['linkedin_url']}")

            contact_info = profile_data.get('contact_info', {})
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])
            if contact_info.get('github'):
                contact_parts.append(f"GitHub: {contact_info['github']}")

            if contact_parts:
                contact_text = ' | '.join(contact_parts)
                content.append(Paragraph(contact_text, style_contact))

            content.append(Spacer(1, 0.1*inch))

            # Professional Summary (if derivable from work history)
            work_history = profile_data.get('work_history', [])
            if work_history:
                content.append(Paragraph('PROFESSIONAL SUMMARY', style_section))
                # Generate summary based on most recent role
                recent_job = work_history[0] if work_history else {}
                title = recent_job.get('title', 'Professional')
                summary = f"Experienced {title} with proven track record in technology and innovation. "
                summary += "Specialized in cloud engineering, DevOps, and AI technologies."
                content.append(Paragraph(summary, style_body))
                content.append(Spacer(1, 0.15*inch))

            # Education Section
            education = profile_data.get('education', [])
            if education:
                content.append(Paragraph('EDUCATION', style_section))
                for edu in education:
                    degree = edu.get('degree', 'Degree')
                    institution = edu.get('institution', 'Institution')
                    year = edu.get('year', 'Year')

                    content.append(Paragraph(f"<b>{degree}</b>", style_subsection))
                    content.append(Paragraph(f"{institution} | {year}", style_date))
                    content.append(Spacer(1, 0.1*inch))

            # Work Experience Section
            if work_history:
                content.append(Paragraph('PROFESSIONAL EXPERIENCE', style_section))
                for job in work_history:
                    title = job.get('title', 'Position')
                    company = job.get('company', 'Company')
                    dates = self.format_date_range(job.get('dates', ''))

                    content.append(Paragraph(f"<b>{title}</b>", style_subsection))
                    content.append(Paragraph(f"{company} | {dates}", style_date))

                    bullets = job.get('bullets', [])
                    for bullet in bullets:
                        bullet_text = f"• {bullet}"
                        content.append(Paragraph(bullet_text, style_body))

                    content.append(Spacer(1, 0.1*inch))

            # Projects Section
            projects = profile_data.get('projects', [])
            if projects:
                content.append(Paragraph('PROJECTS & RESEARCH', style_section))
                for project in projects:
                    name = project.get('name', 'Project')
                    description = project.get('description', '')
                    technologies = project.get('technologies', [])
                    link = project.get('link', '')

                    content.append(Paragraph(f"<b>{name}</b>", style_subsection))
                    if description:
                        content.append(Paragraph(description, style_body))
                    if technologies:
                        tech_text = f"Technologies: {', '.join(technologies)}"
                        content.append(Paragraph(tech_text, style_date))
                    if link:
                        content.append(Paragraph(f"Link: {link}", style_date))
                    content.append(Spacer(1, 0.1*inch))

            # Skills Section
            skills = profile_data.get('skills', [])
            if skills:
                content.append(Paragraph('TECHNICAL SKILLS', style_section))
                # Group skills for better presentation
                skills_text = ', '.join(skills)
                content.append(Paragraph(skills_text, style_body))
                content.append(Spacer(1, 0.15*inch))

            # Certifications (if available in contact_info or custom field)
            certifications = profile_data.get('certifications', [])
            if certifications:
                content.append(Paragraph('CERTIFICATIONS', style_section))
                for cert in certifications:
                    content.append(Paragraph(f"• {cert}", style_body))
                content.append(Spacer(1, 0.15*inch))

            # Footer with generation date
            footer_text = f"CV Generated: {datetime.now().strftime('%B %Y')}"
            content.append(Spacer(1, 0.2*inch))
            content.append(Paragraph(footer_text, style_date))

            # Build PDF
            doc.build(content)

            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()

            return True, pdf_bytes, None

        except Exception as e:
            return False, None, f"Error generating CV PDF: {str(e)}"

    def generate_docx(self, profile_data: Dict[str, Any]) -> Tuple[bool, Optional[bytes], Optional[str]]:
        """
        Generate comprehensive CV in DOCX format

        Returns:
            (success, docx_bytes, error_message)
        """
        try:
            # Create document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Header - Name
            name = profile_data.get('name', 'Your Name')
            heading = doc.add_heading(name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.size = Pt(20)
            heading.runs[0].font.color.rgb = RGBColor(10, 37, 64)

            # Contact Information
            contact_parts = []
            if profile_data.get('email'):
                contact_parts.append(profile_data['email'])
            if profile_data.get('phone'):
                contact_parts.append(profile_data['phone'])
            if profile_data.get('linkedin_url'):
                contact_parts.append(f"LinkedIn: {profile_data['linkedin_url']}")

            contact_info = profile_data.get('contact_info', {})
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])
            if contact_info.get('github'):
                contact_parts.append(f"GitHub: {contact_info['github']}")

            if contact_parts:
                contact_text = ' | '.join(contact_parts)
                contact_para = doc.add_paragraph(contact_text)
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_para.runs[0].font.size = Pt(10)
                contact_para.runs[0].font.color.rgb = RGBColor(51, 51, 51)

            doc.add_paragraph()  # Spacing

            # Professional Summary
            work_history = profile_data.get('work_history', [])
            if work_history:
                doc.add_heading('PROFESSIONAL SUMMARY', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                recent_job = work_history[0] if work_history else {}
                title = recent_job.get('title', 'Professional')
                summary = f"Experienced {title} with proven track record in technology and innovation. "
                summary += "Specialized in cloud engineering, DevOps, and AI technologies."
                doc.add_paragraph(summary)

            # Education Section
            education = profile_data.get('education', [])
            if education:
                doc.add_heading('EDUCATION', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                for edu in education:
                    degree = edu.get('degree', 'Degree')
                    institution = edu.get('institution', 'Institution')
                    year = edu.get('year', 'Year')

                    p = doc.add_paragraph()
                    p.add_run(degree).bold = True
                    doc.add_paragraph(f"{institution} | {year}").runs[0].font.italic = True

            # Work Experience Section
            if work_history:
                doc.add_heading('PROFESSIONAL EXPERIENCE', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                for job in work_history:
                    title = job.get('title', 'Position')
                    company = job.get('company', 'Company')
                    dates = self.format_date_range(job.get('dates', ''))

                    p = doc.add_paragraph()
                    p.add_run(title).bold = True
                    doc.add_paragraph(f"{company} | {dates}").runs[0].font.italic = True

                    bullets = job.get('bullets', [])
                    for bullet in bullets:
                        doc.add_paragraph(bullet, style='List Bullet')

            # Projects Section
            projects = profile_data.get('projects', [])
            if projects:
                doc.add_heading('PROJECTS & RESEARCH', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                for project in projects:
                    name = project.get('name', 'Project')
                    description = project.get('description', '')
                    technologies = project.get('technologies', [])
                    link = project.get('link', '')

                    p = doc.add_paragraph()
                    p.add_run(name).bold = True
                    if description:
                        doc.add_paragraph(description)
                    if technologies:
                        tech_text = f"Technologies: {', '.join(technologies)}"
                        doc.add_paragraph(tech_text).runs[0].font.italic = True
                    if link:
                        doc.add_paragraph(f"Link: {link}").runs[0].font.size = Pt(9)

            # Skills Section
            skills = profile_data.get('skills', [])
            if skills:
                doc.add_heading('TECHNICAL SKILLS', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                skills_text = ', '.join(skills)
                doc.add_paragraph(skills_text)

            # Certifications
            certifications = profile_data.get('certifications', [])
            if certifications:
                doc.add_heading('CERTIFICATIONS', level=2).runs[0].font.color.rgb = RGBColor(10, 37, 64)
                for cert in certifications:
                    doc.add_paragraph(cert, style='List Bullet')

            # Footer with generation date
            footer_text = f"CV Generated: {datetime.now().strftime('%B %Y')}"
            footer_para = doc.add_paragraph(footer_text)
            footer_para.runs[0].font.size = Pt(9)
            footer_para.runs[0].font.italic = True
            footer_para.runs[0].font.color.rgb = RGBColor(102, 102, 102)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            return True, docx_bytes, None

        except Exception as e:
            return False, None, f"Error generating CV DOCX: {str(e)}"

    def generate_cv_files(self, profile_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate both PDF and DOCX CV files

        Returns:
            {
                "success": bool,
                "pdf_bytes": bytes or None,
                "docx_bytes": bytes or None,
                "error": str or None
            }
        """
        # Generate PDF
        pdf_success, pdf_bytes, pdf_error = self.generate_pdf(profile_data)

        # Generate DOCX
        docx_success, docx_bytes, docx_error = self.generate_docx(profile_data)

        if pdf_success and docx_success:
            return {
                "success": True,
                "pdf_bytes": pdf_bytes,
                "docx_bytes": docx_bytes,
                "error": None
            }
        elif pdf_success:
            return {
                "success": True,
                "pdf_bytes": pdf_bytes,
                "docx_bytes": None,
                "error": f"DOCX generation failed: {docx_error}"
            }
        elif docx_success:
            return {
                "success": True,
                "pdf_bytes": None,
                "docx_bytes": docx_bytes,
                "error": f"PDF generation failed: {pdf_error}"
            }
        else:
            return {
                "success": False,
                "pdf_bytes": None,
                "docx_bytes": None,
                "error": f"Both generations failed. PDF: {pdf_error}, DOCX: {docx_error}"
            }


# Singleton instance
_cv_generator = None


def get_cv_generator() -> CVGenerator:
    """Get or create CV generator singleton"""
    global _cv_generator
    if _cv_generator is None:
        _cv_generator = CVGenerator()
    return _cv_generator
