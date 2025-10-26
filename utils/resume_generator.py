"""
Resume generator - Create ATS-optimized PDF and DOCX files
"""

import io
from typing import Dict, Any, Optional, Tuple, List
from datetime import datetime
import streamlit as st

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# DOCX generation
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ResumeGenerator:
    """Generate ATS-optimized résumés in PDF and DOCX formats"""

    def __init__(self):
        pass

    def format_date_range(self, dates: str) -> str:
        """Format date range for consistency"""
        if not dates:
            return ""
        # Basic cleanup
        dates = dates.strip()
        dates = dates.replace(" - ", " to ")
        return dates

    def generate_pdf(self, profile_data: Dict[str, Any]) -> Tuple[bool, Optional[bytes], Optional[str]]:
        """
        Generate ATS-friendly PDF résumé

        Returns:
            (success, pdf_bytes, error_message)
        """
        try:
            # Create PDF buffer
            buffer = io.BytesIO()

            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            # Styles
            styles = getSampleStyleSheet()

            # Custom styles for ATS compatibility (simple, no complex formatting)
            style_name = ParagraphStyle(
                'CustomName',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=6,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )

            style_contact = ParagraphStyle(
                'CustomContact',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=12,
                alignment=TA_CENTER
            )

            style_heading = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=12,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=6,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            )

            style_subheading = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=4,
                fontName='Helvetica-Bold'
            )

            style_body = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#333333'),
                spaceAfter=4,
                leftIndent=20
            )

            # Build content
            story = []

            # Name
            name = profile_data.get('name', 'Your Name')
            story.append(Paragraph(name, style_name))

            # Contact info
            contact_parts = []
            if profile_data.get('email'):
                contact_parts.append(profile_data['email'])
            if profile_data.get('phone'):
                contact_parts.append(profile_data['phone'])
            if profile_data.get('linkedin_url'):
                contact_parts.append(profile_data['linkedin_url'])

            contact_info = profile_data.get('contact_info', {})
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])

            if contact_parts:
                story.append(Paragraph(" | ".join(contact_parts), style_contact))

            story.append(Spacer(1, 0.2*inch))

            # Work Experience
            work_history = profile_data.get('work_history', [])
            if work_history:
                story.append(Paragraph("WORK EXPERIENCE", style_heading))
                story.append(Spacer(1, 0.1*inch))

                for job in work_history:
                    # Job title and company
                    job_line = f"<b>{job.get('title', '')}</b> | {job.get('company', '')}"
                    story.append(Paragraph(job_line, style_subheading))

                    # Dates
                    dates = self.format_date_range(job.get('dates', ''))
                    if dates:
                        story.append(Paragraph(dates, style_body))

                    # Bullets
                    bullets = job.get('bullets', [])
                    for bullet in bullets[:5]:  # Limit to 5 bullets
                        story.append(Paragraph(f"• {bullet}", style_body))

                    story.append(Spacer(1, 0.15*inch))

            # Skills
            skills = profile_data.get('skills', [])
            if skills:
                story.append(Paragraph("SKILLS", style_heading))
                skills_text = ", ".join(skills)
                story.append(Paragraph(skills_text, style_body))
                story.append(Spacer(1, 0.15*inch))

            # Projects
            projects = profile_data.get('projects', [])
            if projects:
                story.append(Paragraph("PROJECTS", style_heading))
                story.append(Spacer(1, 0.1*inch))

                for project in projects:
                    # Project name
                    project_name = project.get('name', '')
                    story.append(Paragraph(f"<b>{project_name}</b>", style_subheading))

                    # Description
                    desc = project.get('description', '')
                    if desc:
                        story.append(Paragraph(desc, style_body))

                    # Technologies
                    techs = project.get('technologies', [])
                    if techs:
                        tech_text = f"Technologies: {', '.join(techs)}"
                        story.append(Paragraph(tech_text, style_body))

                    story.append(Spacer(1, 0.15*inch))

            # Education
            education = profile_data.get('education', [])
            if education:
                story.append(Paragraph("EDUCATION", style_heading))
                story.append(Spacer(1, 0.1*inch))

                for edu in education:
                    edu_line = f"<b>{edu.get('degree', '')}</b> | {edu.get('institution', '')}"
                    story.append(Paragraph(edu_line, style_subheading))

                    year = edu.get('year', '')
                    if year:
                        story.append(Paragraph(year, style_body))

                    story.append(Spacer(1, 0.1*inch))

            # Build PDF
            doc.build(story)

            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()

            return True, pdf_bytes, None

        except Exception as e:
            return False, None, f"Error generating PDF: {str(e)}"

    def generate_docx(self, profile_data: Dict[str, Any]) -> Tuple[bool, Optional[bytes], Optional[str]]:
        """
        Generate ATS-friendly DOCX résumé

        Returns:
            (success, docx_bytes, error_message)
        """
        try:
            # Create document
            doc = Document()

            # Set margins (1 inch all around for ATS compatibility)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Name
            name = profile_data.get('name', 'Your Name')
            heading = doc.add_heading(name, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info
            contact_parts = []
            if profile_data.get('email'):
                contact_parts.append(profile_data['email'])
            if profile_data.get('phone'):
                contact_parts.append(profile_data['phone'])
            if profile_data.get('linkedin_url'):
                contact_parts.append(profile_data['linkedin_url'])

            if contact_parts:
                contact = doc.add_paragraph(" | ".join(contact_parts))
                contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in contact.runs:
                    run.font.size = Pt(10)

            doc.add_paragraph()  # Spacer

            # Work Experience
            work_history = profile_data.get('work_history', [])
            if work_history:
                doc.add_heading('WORK EXPERIENCE', level=2)

                for job in work_history:
                    # Job title and company
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{job.get('title', '')} | {job.get('company', '')}")
                    job_run.bold = True
                    job_run.font.size = Pt(11)

                    # Dates
                    dates = self.format_date_range(job.get('dates', ''))
                    if dates:
                        date_para = doc.add_paragraph(dates)
                        date_para.style = 'Normal'
                        for run in date_para.runs:
                            run.font.size = Pt(10)

                    # Bullets
                    bullets = job.get('bullets', [])
                    for bullet in bullets[:5]:
                        bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                        for run in bullet_para.runs:
                            run.font.size = Pt(10)

                    doc.add_paragraph()  # Spacer

            # Skills
            skills = profile_data.get('skills', [])
            if skills:
                doc.add_heading('SKILLS', level=2)
                skills_text = ", ".join(skills)
                skills_para = doc.add_paragraph(skills_text)
                for run in skills_para.runs:
                    run.font.size = Pt(10)
                doc.add_paragraph()

            # Projects
            projects = profile_data.get('projects', [])
            if projects:
                doc.add_heading('PROJECTS', level=2)

                for project in projects:
                    # Project name
                    project_para = doc.add_paragraph()
                    project_run = project_para.add_run(project.get('name', ''))
                    project_run.bold = True
                    project_run.font.size = Pt(11)

                    # Description
                    desc = project.get('description', '')
                    if desc:
                        desc_para = doc.add_paragraph(desc)
                        for run in desc_para.runs:
                            run.font.size = Pt(10)

                    # Technologies
                    techs = project.get('technologies', [])
                    if techs:
                        tech_text = f"Technologies: {', '.join(techs)}"
                        tech_para = doc.add_paragraph(tech_text)
                        for run in tech_para.runs:
                            run.font.size = Pt(10)

                    doc.add_paragraph()

            # Education
            education = profile_data.get('education', [])
            if education:
                doc.add_heading('EDUCATION', level=2)

                for edu in education:
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{edu.get('degree', '')} | {edu.get('institution', '')}")
                    edu_run.bold = True
                    edu_run.font.size = Pt(11)

                    year = edu.get('year', '')
                    if year:
                        year_para = doc.add_paragraph(year)
                        for run in year_para.runs:
                            run.font.size = Pt(10)

                    doc.add_paragraph()

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            return True, docx_bytes, None

        except Exception as e:
            return False, None, f"Error generating DOCX: {str(e)}"

    def generate_resume_files(self, profile_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate both PDF and DOCX résumés

        Returns:
            {
                "success": bool,
                "pdf_bytes": bytes or None,
                "docx_bytes": bytes or None,
                "errors": List[str]
            }
        """
        errors = []

        # Generate PDF
        pdf_success, pdf_bytes, pdf_error = self.generate_pdf(profile_data)
        if not pdf_success:
            errors.append(f"PDF: {pdf_error}")

        # Generate DOCX
        docx_success, docx_bytes, docx_error = self.generate_docx(profile_data)
        if not docx_success:
            errors.append(f"DOCX: {docx_error}")

        return {
            "success": pdf_success or docx_success,  # Success if at least one works
            "pdf_bytes": pdf_bytes if pdf_success else None,
            "docx_bytes": docx_bytes if docx_success else None,
            "errors": errors
        }


# Singleton instance
_resume_generator = None


def get_resume_generator() -> ResumeGenerator:
    """Get or create resume generator singleton"""
    global _resume_generator
    if _resume_generator is None:
        _resume_generator = ResumeGenerator()
    return _resume_generator
